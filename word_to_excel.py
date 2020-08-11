import pandas as pd
import docx
import re
import os


def submission_get(str_submit):# 通过分析word文件第一行的数据得到提交人 提交部门 提交日期这样的信息
    ele_submit_department=""
    ele_submit_people=""
    ele_submit_time=""
    pid=0
    while pid<len(str_submit):# 逐字分析,这部分可以改成正则
        if str_submit[pid]=="提" and str_submit[pid+1]=="交" and str_submit[pid+2]=="部" and str_submit[pid+3]=="门":
            pid+=4
            while str_submit[pid]==":" or str_submit[pid]=="：" or str_submit[pid]==" ":#去除例如提交部门:      abc中abc前面的空格
                pid+=1
            while str_submit[pid]!=" ":
                if(str_submit[pid:pid+2]=="提交"):# 如果提交部门是空的没有填,那么就直接跳出来,下面同理
                    break
                ele_submit_department+=str_submit[pid]
                pid+=1
        elif str_submit[pid]=="提" and str_submit[pid+1]=="交" and str_submit[pid+2]=="人":
            pid+=3
            while str_submit[pid]==":" or str_submit[pid]=="：" or str_submit[pid]==" ":
                pid+=1
            while str_submit[pid]!=" ":
                if(str_submit[pid:pid+2]=="提交"):
                    break
                ele_submit_people+=str_submit[pid]
                pid+=1
        elif str_submit[pid]=="提" and str_submit[pid+1]=="交" and str_submit[pid+2]=="日" and str_submit[pid+3]=="期":
            pid+=4
            while pid<pid<len(str_submit) and str_submit[pid]==":" or str_submit[pid]=="：" or str_submit[pid]==" ":
                pid+=1
            while pid<len(str_submit) and (str_submit[pid]!=" ") :# 防止超过这一行的长度
                if(str_submit[pid:pid+2]=="提交"):
                    break
                ele_submit_time+=str_submit[pid]
                pid+=1
        else:
            pid+=1
    return [ele_submit_department,ele_submit_people,ele_submit_time]
def is_number(s): # 判断是不是数字
    try:
        float(s)
        return True
    except ValueError:
        pass
 
    try:
        import unicodedata
        unicodedata.numeric(s)
        return True
    except (TypeError, ValueError):
        pass
 
    return False
def docx_file_analyze(file,filename):# 分析word文件
    print("段落数:"+str(len(file.paragraphs)))#段落数为13，每个回车隔离一段
    para=file.paragraphs
    #print(para[1].text)
    str_submit=para[1].text # 提取paragraph 的第一行

    ele_submit_department,ele_submit_people,ele_submit_time=submission_get(str_submit)

    #已经提取了提交部门,提交人,提交日期;

    #接下来提取名称,品牌,配置信息,数量,使用人,金额(单价),

    print("表格数:"+str(len(file.tables)))
    table=file.tables[0]

    table_length=len(table.row_cells(0))# 第一行的长度
    table_width=len(table.column_cells(0)) #第一列的长度

    # 下面对表格第一行进行扫描,确定名称,品牌,配置信息,数量,使用人,金额(单价)所在的表项,然后添加之前的提交部门,提交人,提交日期,文件名,作为一个元组添加到list里面
    file_res=[]

    file_map={"名称":-1,"品牌":-1,"配置信息":-1,"数量":-1,"使用人":-1,"单价":-1}# 如果扫描第一行之后还是-1,表示该项缺省;将缺省项添加到备注里面;
    pid=0
    for cell in table.row_cells(0):
        name=cell.text
        #遍历字典
        for k,v in file_map.items():
            if k==name:
                file_map[k]=pid
        pid+=1
    # for k,v in file_map.items():
    #     print(str(k)+":"+str(v))
    # 此时通过file_map可以获得对应信息的列号;
    note_list=[] #缺省项
    for k,v in file_map.items():
        if v==-1:
            note_list.append(k)


    test_col1=file_map["名称"]


    #如果test_col1为空格,说明这个项是空项;
    # test_string="      \n+   "
    # if(test_string.isspace()):
    #     print("yes")
    # else:
    #     print("no")

    for i in range(0,table_width):

        if is_number(table.cell(i,0).text):# 表示和第一行的长度相同,也就是有效的
            if not table.cell(i,test_col1).text.isspace() and table.cell(i,test_col1).text is not None and table.cell(i,test_col1).text!='':# 如果名称那一列不是空格
     
                text_name=" "
                text_brand=" "
                text_infor=" "
                text_num=" "
                text_user=" "
                text_price=" "
                if(file_map["名称"]!=-1): #!=-1表示该表有这个表项,没有的话那就是初始的空格
                    text_name=table.cell(i,file_map["名称"]).text
                if(file_map["品牌"]!=-1):
                    text_brand=table.cell(i,file_map["品牌"]).text
                if(file_map["配置信息"]!=-1):
                    text_infor=table.cell(i,file_map["配置信息"]).text
                if(file_map["数量"]!=-1):
                    text_num=table.cell(i,file_map["数量"]).text
                if(file_map["使用人"]!=-1):
                    text_user=table.cell(i,file_map["使用人"]).text
                if(file_map["单价"]!=-1):
                    text_price=table.cell(i,file_map["单价"]).text
                if text_name==text_price or text_brand==text_price:
                    continue
                file_res.append([text_name,text_brand,text_infor,text_num,text_user,text_price,ele_submit_department,ele_submit_people,ele_submit_time,filename])
                
                
                
            else:
                continue

    return file_res

# 判断文件后缀
def endWith(s, *endstring):
    array = map(s.endswith, endstring)
    if True in array:
        return True
    else:
        return False

def write_into_excel(final_res):# 将数据写进excel,名称是result.xlsx，如果之前有这个文件请删除之后再使用
    excel_file_path=path1+"/result_of_excel.xlsx"
    if os.path.exists(excel_file_path):# 如果有之前的版本,那么删除
        os.remove(excel_file_path)
    if not os.path.exists(excel_file_path):
        df=pd.DataFrame()
        df.to_excel(excel_file_path)
        # 如果没有该excel那么就新建一个;
    #名称,品牌,配置信息,数量,使用人,金额(单价)所在的表项,然后添加之前的提交部门,提交人,提交日期,文件名
    df=pd.read_excel(excel_file_path)
    # 初始化列的数据
    df["名称"]=None
    df["品牌"]=None
    df["配置信息"]=None
    df["数量"]=None
    df["使用人"]=None
    df["金额(单价)"]=None
    df["提交部门"]=None
    df["提交人"]=None
    df["提交日期"]=None
    df["文件名"]=None

    pid=0
    for ele in final_res:# 把final_Res的数据添加到df中以便写入excel
        df.loc[pid]=ele
        pid+=1
    df.to_excel(excel_file_path,index=None)
    print("写入excel文件:"+excel_file_path)



###### 完成部分

final_res=[] # 存储从word文件中分析得来的结果
path1=os.path.abspath('.')
des_folder=input("告诉我文件夹名称:")
log_d=path1+"/"+des_folder # 记录文件夹地址
logfiles=os.listdir(log_d)

for filename in logfiles:# 遍历文件夹里面的文件
    if endWith(filename,'.docx'):
        # 建立文件绝对路径:
        filepath=log_d+'/'+filename
        #读取文件:
        file=docx.Document(filepath)
        temp_res=docx_file_analyze(file,filename)
        for temp in temp_res:
            final_res.append(temp)
        print(filename+" "+"扫描完成")

write_into_excel(final_res)

##### 完成部分如上


##### 测试部分如下:有问题的文件可以使用测试部分进行测试,输出到控制台上;
# path1=os.path.abspath('.')
# file=docx.Document(path1+"/test/信息化设备领用申请单（办公室20191227）.docx")

# temp_res=docx_file_analyze(file,"信息化设备领用申请单（办公室20191227）.docx")
# print(temp_res)

